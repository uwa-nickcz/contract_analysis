# -*- coding: utf-8 -*-
# @Time    : 2025/4/8 16:07
# @Author  : cz
# @File    : data_loader.py
# @Software: PyCharm
import json
from docx import Document
import os
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_core.documents import Document as LangDocument


def get_files(directory):
    res = []
    for root, dirs, files in os.walk(directory):
        for file in files:
            res.append(os.path.join(root, file))
    return res




class DataLoader:

    @staticmethod
    def load_json(path):
        try:
            with open(path, 'r',encoding='utf-8') as f:
                data = json.load(f)
            return data
        except UnicodeDecodeError:
            with open(path, 'r',encoding='gbk') as f:
                data = json.load(f)
            return data




class ContractParser:
    def extract_document_title(self, doc_path):

        file_title = os.path.splitext(os.path.basename(doc_path))[0]
        return file_title

    def parse_document_by_structured(self, doc_path):
        """解析带多级标题的文档结构"""
        doc_title = self.extract_document_title(doc_path)
        doc = Document(doc_path)

        structured_data = []
        current_hierarchy = []

        for para in doc.paragraphs:
            # 检测标题样式（假设使用Word的标题样式）
            if para.style.name.startswith('Heading'):
                level = int(para.style.name[-1])
                text = para.text.strip()

                # 维护标题层级
                current_hierarchy = current_hierarchy[:level-1] + [text]

            # 处理表格
            elif para._element.tag.endswith('tbl'):
                table = para._element
                table_data = [
                    [cell.text.strip() for cell in row.cells]
                    for row in table.rows
                ]
                structured_data.append({
                    "type": "table",
                    "content": table_data,
                    "hierarchy": current_hierarchy.copy(),
                    "doc_title": doc_title
                })

            # 处理普通段落
            elif para.text.strip():
                structured_data.append({
                    "type": "text",
                    "content": para.text.strip(),
                    "hierarchy": current_hierarchy.copy(),
                    "doc_title": doc_title
                })

        return structured_data

    def parse_document_by_chunk(self, doc_path,chunk_size=512,chunk_overlap=128):
        """解析文档为大小为chunk_size的块"""
        doc_title = self.extract_document_title(doc_path)
        doc = Document(doc_path)

        txt_data = []
        pure_txt_data = []
        current_hierarchy = []

        for para in doc.paragraphs:
            txt_data.append(para.text.strip())

        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:  # 忽略空单元格
                        row_text.append(cell_text)  # 将单元格内容加入纯文本数据
                if row_text:  # 如果该行有内容
                    txt_data.append("\t".join(row_text))

        full_text = "\n".join(txt_data)


        pure_txt_data.append({
            "type": "text",
            "content": full_text,
            "hierarchy": current_hierarchy.copy(),
            "doc_title": doc_title
        })

        return pure_txt_data


class ContractChunker:
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=512,
            chunk_overlap=128,
            separators=["\n\n", "\n", "。", "！", "？"]
        )

    def generate_chunks(self, structured_data):
        chunks = []
        for item in structured_data:
            # 构建多级元数据
            metadata = {
                "doc_title": item["doc_title"],
                "section_hierarchy": " > ".join(item["hierarchy"]),
                "content_type": item["type"]
            }

            if item["type"] == "table":
                # 表格转换为Markdown格式
                markdown_table = "\n".join(
                    "| " + " | ".join(row) + " |"
                    for row in item["content"]
                )
                chunks.append(LangDocument(
                    page_content=f"表格内容：\n{markdown_table}",
                    metadata=metadata
                ))
            else:
                # 文本分块增强上下文
                prefix = f"文档标题：{metadata['doc_title']}\n章节路径：{metadata['section_hierarchy']}\n"
                splits = self.text_splitter.split_text(prefix + item["content"])

                for split in splits:
                    chunks.append(LangDocument(
                        page_content=split,
                        metadata=metadata
                    ))
        return chunks